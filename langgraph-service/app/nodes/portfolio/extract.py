"""
Portfolio extraction — parses an uploaded resume/activity list/course list
(PDF or DOCX) into structured portfolio pieces via one Haiku call.

The extracted list is returned to the frontend for review/edit/selection;
nothing is persisted here. Saving happens client-side (direct Supabase insert)
after the user confirms in the review modal.
"""
import io
import json
import logging
import re

from anthropic import AsyncAnthropic

from app.config import ANTHROPIC_API_KEY

logger = logging.getLogger(__name__)

_anthropic = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
HAIKU = "claude-haiku-4-5-20251001"

CATEGORIES = ["experience", "volunteering", "award", "course", "certification", "club", "skill", "other"]

MAX_FILE_BYTES = 5 * 1024 * 1024   # 5MB — resumes are tiny; anything bigger is wrong
MAX_TEXT_CHARS = 20_000            # bounds the Haiku input for cost
MAX_ITEMS      = 40


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx_text(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_file_text(filename: str, content: bytes) -> str:
    """Raw text from a PDF/DOCX upload. Raises ValueError with a user-facing message."""
    if len(content) > MAX_FILE_BYTES:
        raise ValueError("File is too large. Please upload a file under 5MB.")
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            text = _extract_pdf_text(content)
        elif name.endswith(".docx"):
            text = _extract_docx_text(content)
        else:
            raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")
    except ValueError:
        raise
    except Exception as exc:
        logger.warning(f"[portfolio] file parse failed for {filename!r}: {exc}")
        raise ValueError("Could not read that file. Make sure it is a valid PDF or DOCX.")
    text = re.sub(r"[ \t]+", " ", text).strip()
    if len(text) < 40:
        raise ValueError("Could not find readable text in that file. Scanned images are not supported.")
    return text[:MAX_TEXT_CHARS]


def _parse_items(text: str):
    """Permissive JSON parse: direct, then first [...] block. Returns list or None."""
    for candidate in (text, *(m.group(0) for m in [re.search(r"\[[\s\S]*\]", text)] if m)):
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, list):
                return parsed
            if isinstance(parsed, dict) and isinstance(parsed.get("items"), list):
                return parsed["items"]
        except Exception:
            continue
    return None


EXTRACTION_PROMPT = """You extract structured portfolio pieces from a student's resume, activity list, brag sheet, or course list.

Pull out every distinct item: jobs, internships, volunteer work, awards, honors, courses (AP/IB/dual enrollment/online), certifications, clubs, leadership roles, notable skills, and projects.

Rules:
- category must be exactly one of: experience, volunteering, award, course, certification, club, skill, other
- title: short and specific, max 80 characters (e.g. "Software Engineering Intern at Acme", "AP Computer Science A", "DECA State Finalist")
- description: 1-2 sentences capturing the concrete details present in the document (dates, role, scope, results). Use only what the document says, do not invent details. Empty string if the document gives nothing beyond the title.
- Skip contact info, objective/summary paragraphs, and references.
- Never use em dashes anywhere. Use commas or periods instead.
- Return at most {max_items} items, the most substantive ones.

Return ONLY a valid JSON array, no other text, no markdown, no backticks:
[{{"category": "...", "title": "...", "description": "..."}}]

DOCUMENT:
{document}"""


async def extract_portfolio_items(document: str) -> list[dict]:
    """
    Extracted document text -> list of {category, title, description}.
    Raises ValueError with a user-facing message on expected failures.
    (File reading/validation lives in extract_file_text so the endpoint can
    reject bad files BEFORE burning a rate-limited upload.)
    """
    resp = await _anthropic.messages.create(
        model=HAIKU,
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": EXTRACTION_PROMPT.format(max_items=MAX_ITEMS, document=document),
        }],
    )
    raw = resp.content[0].text if resp.content else ""
    if resp.stop_reason == "max_tokens":
        logger.warning("[portfolio] extraction truncated at max_tokens")
    parsed = _parse_items(raw)
    if parsed is None:
        logger.warning(f"[portfolio] extraction parse failed: {raw[:200]!r}")
        raise ValueError("Could not extract items from that file. Try again or add pieces manually.")

    items = []
    for entry in parsed[:MAX_ITEMS]:
        if not isinstance(entry, dict):
            continue
        title = str(entry.get("title") or "").strip()[:120]
        if not title:
            continue
        category = str(entry.get("category") or "").strip().lower()
        if category not in CATEGORIES:
            category = "other"
        description = str(entry.get("description") or "").strip()[:500]
        items.append({"category": category, "title": title, "description": description})

    if not items:
        raise ValueError("No portfolio items were found in that file. Try adding pieces manually.")

    logger.info(f"[portfolio] extracted {len(items)} items")
    return items
